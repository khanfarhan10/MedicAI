# Implement Two Functionalities, Replace Text & Replace Image
import fnmatch
import os
import zipfile
import shutil
from docx import Document

def ReplaceText(InputFilepath, OutputFilepath, ReplacementDictionary):
    document = Document(InputFilepath)
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            print(text)
            if text in ReplacementDictionary.keys():
                text = text.replace(text, ReplacementDictionary[text])
                inline[i].text = text

    document.save(OutputFilepath)

def find(pattern, path):
    result = []
    for root, dirs, files in os.walk(path):
        for name in files:
            if fnmatch.fnmatch(name, pattern):
                result.append(os.path.join(root, name))
    return result

def ExtractZipFile(SourceZipFile,OutputDirectory):
    with zipfile.ZipFile(SourceZipFile, 'r') as zip_ref:
        zip_ref.extractall(OutputDirectory)

def ReplaceWordImage(SrcImgFilePath,DestImgFileName,OutputDirectory):
    DestImgFilePath=os.path.join(OutputDirectory, "word", "media",DestImgFileName)
    # First Delete the Destination File & Then Copy the SourceFile to the DestinationFile
    os.remove(DestImgFilePath)
    shutil.copy(SrcImgFilePath, DestImgFilePath)

def zip_directory(folder_path, zip_path):
    with zipfile.ZipFile(zip_path, mode='w') as zipf:
        len_dir_path = len(folder_path)
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, file_path[len_dir_path:])

def ReconsructDocument(OutputDirectory,OutputFilepath):
    zip_directory(OutputDirectory,OutputFilepath)
    
def create_dir(dir):
  if not os.path.exists(dir):
    os.makedirs(dir)
    # print("Created Directory : ", dir)
    return dir
if __name__== "__main__":
    # code here
    InputFilepath = os.path.join('CertificateTemplates','Government_Certificate_Sample_Unreal.docx')
    IntermediateDirectory='IntermediateFiles'
    create_dir(IntermediateDirectory)
    IntermediateFilepath = os.path.join(IntermediateDirectory,'Government_Certificate_Generated_Intermediate.docx')
    OutputDirec='Outputs'
    create_dir(OutputDirec)
    OutputFilepath = os.path.join(OutputDirec,'Government_Certificate_Generated_Final.docx')
    OutputDirectory='/TemporaryFiles/'
    create_dir(OutputDirectory)
    SrcImgFilePath='BarCodeImages/ReplacementImage.png'
    DestImgFileName='image4.png'
    

    ReplacementDictionary = {'Your_Amazing_Name': 'Farhan Hai Khan',
                            'YOUR_AWESOME_NAME_HERE': 'FARHAN HAI KHAN',
                            'COVID': 'Coronavirus Scale : 11.67 %',
                            'Automated': 'Automated Tests : Passed Successfully',
                            'Manual': 'Manual Tests : Passed with Considerations',
                            'Description': 'Comments : Fit for Travel',
                            'UID_Aadhaar': 'UID (Aadhaar) : 7458-2541-7364',
                            'Cert': 'Certificate ID : 9546-8741-9463',
                            'Generated': 'Generated : 22.06.2020 5:30GMT',
                            'Signature': 'Rameshwar Manav',
                            'Esteemed': 'RAMESHWAR MANAV',
                            'Designation': 'Chief Engg. DVC',
                            'Autograph': 'Shreya Shambhavi',
                            'Responsibility': 'SHREYA SHAMBHAVI',
                            'Pos': 'Radiologist RIMS'
                            }
    
    ReplaceText(InputFilepath, IntermediateFilepath, ReplacementDictionary)
    ExtractZipFile(IntermediateFilepath,OutputDirectory)
    
    ReplaceWordImage(SrcImgFilePath,DestImgFileName,OutputDirectory)
    ReconsructDocument(OutputDirectory,OutputFilepath)
    
    print('Docx File Written To :', OutputFilepath)
    
    
# python CertGen.py

"""
src_file = "Government_Certificate_Generated_Unreal.docx"
out_dir = "temp_run/"
des_file = "Government_Certificate_Generated_Unreal_IMG_REPLACED.docx"

with zipfile.ZipFile(src_file, 'r') as zip_ref:
    zip_ref.extractall(out_dir)
    


imgrep_path = "image4.png"
# it should be png format

media_path = os.path.join(out_dir, "word", "media")
# img_file=find('*image4.png', media_path)
img_file = os.path.join(media_path, "image4.png")

os.remove(img_file)


shutil.copy(imgrep_path, img_file)


def zip_directory(folder_path, zip_path):
    with zipfile.ZipFile(zip_path, mode='w') as zipf:
        len_dir_path = len(folder_path)
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, file_path[len_dir_path:])


zip_directory(out_dir, des_file)

"""

"""

document = Document('Government_Certificate_Sample.docx')

dic = {'Your_Amazing_Name': 'Farhan Hai Khan',
       'YOUR_AWESOME_NAME_HERE': 'FARHAN HAI KHAN',
       'COVID': 'Coronavirus Scale : 11.67 %',
       'Automated': 'Automated Tests : Passed Successfully',
       'Manual': 'Manual Tests : Passed with Considerations',
       'Description': 'Comments : Fit for Travel',
       'UID_Aadhaar': 'UID (Aadhaar) : 7458-2541-7364',
       'Cert': 'Certificate ID : 9546-8741-9463',
       'Generated': 'Generated : 22.06.2020 5:30GMT',
       'Signature': 'Rameshwar Manav',
       'Esteemed': 'RAMESHWAR MANAV',
       'Designation': 'Chief Engg. DVC',
       'Autograph': 'Shreya Shambhavi',
       'Responsibility': 'SHREYA SHAMBHAVI',
       'Pos': 'Radiologist RIMS'
       }
for p in document.paragraphs:
    inline = p.runs
    for i in range(len(inline)):
        text = inline[i].text
        print(text)
        if text in dic.keys():
            text = text.replace(text, dic[text])
            inline[i].text = text

document.save('Government_Certificate_Generated.docx')

# python ImageReplacer.py


def find(pattern, path):
    result = []
    for root, dirs, files in os.walk(path):
        for name in files:
            if fnmatch.fnmatch(name, pattern):
                result.append(os.path.join(root, name))
    return result


src_file = "Government_Certificate_Generated_Unreal.docx"
out_dir = "temp_run/"
des_file = "Government_Certificate_Generated_Unreal_IMG_REPLACED.docx"

with zipfile.ZipFile(src_file, 'r') as zip_ref:
    zip_ref.extractall(out_dir)

imgrep_path = "image4.png"
# it should be png format

media_path = os.path.join(out_dir, "word", "media")
# img_file=find('*image4.png', media_path)
img_file = os.path.join(media_path, "image4.png")

os.remove(img_file)


shutil.copy(imgrep_path, img_file)


def zip_directory(folder_path, zip_path):
    with zipfile.ZipFile(zip_path, mode='w') as zipf:
        len_dir_path = len(folder_path)
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, file_path[len_dir_path:])


zip_directory(out_dir, des_file)


# TODO save to pdf
"""
